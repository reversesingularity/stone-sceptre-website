"""
build_manuscript.py
-------------------
Assembles NephilimChronicles_Book1_MANUSCRIPT.docx from source .md files.
Targets Amazon KDP Paperback: 6" x 9", Georgia 12pt, KDP-compliant margins.

Requirements:
    pip install python-docx
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ---------------------------------------------------------------------------
# CONFIGURATION
# ---------------------------------------------------------------------------

MANUSCRIPT_DIR = Path(os.environ.get(
    "KDP_MANUSCRIPT_DIR",
    r"f:\Projects-cmodi.000\book_writer_ai_toolkit\output\nephilim_chronicles\MANUSCRIPT\book_1"
))
OUTPUT_FILE    = Path(os.environ.get(
    "KDP_OUTPUT_FILE",
    r"f:\Projects-cmodi.000\book_writer_ai_toolkit\output\nephilim_chronicles\NephilimChronicles_Book1_MANUSCRIPT.docx"
))

BODY_FONT      = "Georgia"
BODY_SIZE      = Pt(12)
HEADING_FONT   = "Georgia"
SCENE_BREAK    = "✦"

# KDP 6×9 margins (in inches)
MARGIN_TOP     = Inches(0.75)
MARGIN_BOTTOM  = Inches(0.75)
MARGIN_INSIDE  = Inches(0.875)   # gutter
MARGIN_OUTSIDE = Inches(0.5)

# Source files in assembly order
SOURCE_FILES = [
    MANUSCRIPT_DIR / "BOOK_1_FRONT_MATTER.md",
    MANUSCRIPT_DIR / "BOOK_1_PROLOGUE.md",
    MANUSCRIPT_DIR / "CHAPTER_01_THE_AWAKENING.md",
    MANUSCRIPT_DIR / "CHAPTER_02_THE_HUNTER.md",
    MANUSCRIPT_DIR / "CHAPTER_03_THE_COORDINATES.md",
    MANUSCRIPT_DIR / "CHAPTER_04_THE_DATA_TRAIL.md",
    MANUSCRIPT_DIR / "CHAPTER_05_THE_CELESTIAL_BIOLOGY_LESSON.md",
    MANUSCRIPT_DIR / "CHAPTER_06_GUNS_LOTS_OF_GUNS.md",
    MANUSCRIPT_DIR / "CHAPTER_07_THE_OTHER_HUNTER.md",
    MANUSCRIPT_DIR / "CHAPTER_08_THE_COLLISION.md",
    MANUSCRIPT_DIR / "CHAPTER_09_THE_ARSENAL.md",
    MANUSCRIPT_DIR / "CHAPTER_10_THE_RECKONING.md",
    MANUSCRIPT_DIR / "CHAPTER_11_THE_DEBT_REPAID.md",
    MANUSCRIPT_DIR / "CHAPTER_12_THE_QUESTION.md",
    MANUSCRIPT_DIR / "CHAPTER_13_THE_FREQUENCY.md",
    MANUSCRIPT_DIR / "CHAPTER_14_THE_SILENCE.md",
    MANUSCRIPT_DIR / "CHAPTER_15_THE_DISSOLUTION.md",
    MANUSCRIPT_DIR / "CHAPTER_16_THE_QUEEN_MOVES.md",
    MANUSCRIPT_DIR / "CHAPTER_17_THE_TESTIMONY_OF_THE_ARCHANGEL.md",
    MANUSCRIPT_DIR / "EPILOGUE_THE_DIGGING_BEGINS.md",
    MANUSCRIPT_DIR / "BOOK_1_APPENDICES.md",
]

# ---------------------------------------------------------------------------
# BOOK SELECTION & BOOK 2 CONFIGURATION
# ---------------------------------------------------------------------------

BOOK = int(os.environ.get("KDP_BOOK", "1"))

MANUSCRIPT_DIR_BOOK2 = Path(
    r"f:\Projects-cmodi.000\book_writer_ai_toolkit\output\nephilim_chronicles\MANUSCRIPT\book_2\CHAPTERS"
)
OUTPUT_FILE_BOOK2 = Path(
    r"f:\Projects-cmodi.000\book_writer_ai_toolkit\output\nephilim_chronicles\NephilimChronicles_Book2_MANUSCRIPT.docx"
)

SOURCE_FILES_BOOK2 = [
    MANUSCRIPT_DIR_BOOK2 / "BOOK_2_FRONT_MATTER.md",
    MANUSCRIPT_DIR_BOOK2 / "PROLOGUE_SCENE1_TheFountainsOfTheDeep.md",
    MANUSCRIPT_DIR_BOOK2 / "PROLOGUE_SCENE2_TheTowerAndTheThrone.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_01_TheTwentyNames_REVISED.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_02_DeadReckoning.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_03_TheCaptainsDomain.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_04_TheSameWord.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_05_TheArmarosDomain.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_06_HallowedConflagration.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_07_TheAccidentalFianna.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_07_5_TheCounterHarmonic.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_08_TheGate.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_09_TheLabyrinthOfNight.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_10_EchoesOfJudgment.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_11_TheBreach.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_12_TheJosephiteGambit.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_13_TheArmarosInversion.md",
    MANUSCRIPT_DIR_BOOK2 / "CHAPTER_14_TheAppointedTime.md",
    MANUSCRIPT_DIR_BOOK2 / "EPILOGUE_TheWitnessesInEden.md",
    MANUSCRIPT_DIR_BOOK2 / "BOOK_2_APPENDICES.md",
]

IMAGE_DIR_BOOK2 = Path(
    r"f:\Projects-cmodi.000\book_writer_ai_toolkit\output\nephilim_chronicles\MANUSCRIPT\book_2\IMAGES\kdp_ready\chapters"
)

# Book 3 Configuration
MANUSCRIPT_DIR_BOOK3 = Path(
    r"f:\Projects-cmodi.000\book_writer_ai_toolkit\output\nephilim_chronicles\MANUSCRIPT\book_3\CHAPTERS"
)
OUTPUT_FILE_BOOK3 = Path(
    r"f:\Projects-cmodi.000\book_writer_ai_toolkit\output\nephilim_chronicles\NephilimChronicles_Book3_MANUSCRIPT.docx"
)

SOURCE_FILES_BOOK3 = [
    MANUSCRIPT_DIR_BOOK3 / "prologue.md",
    MANUSCRIPT_DIR_BOOK3 / "CHAPTER_01_TheArchitectureOfResistance.md",
    MANUSCRIPT_DIR_BOOK3 / "CHAPTER_02_TheFrequencyBeneathTheCure.md",
    MANUSCRIPT_DIR_BOOK3 / "CHAPTER_03_TheLastCartography.md",
    MANUSCRIPT_DIR_BOOK3 / "CHAPTER_04_TheGauntlet.md",
    # Add subsequent Book 3 chapters here as they are drafted
]

IMAGE_DIR_BOOK3 = Path(
    r"f:\Projects-cmodi.000\book_writer_ai_toolkit\output\nephilim_chronicles\MANUSCRIPT\book_3\IMAGES\kdp_ready\chapters"
)

# Chapter filename pattern → image filename (order matters: longer prefixes first)
CHAPTER_IMAGE_MAP = {
    "PROLOGUE_SCENE1": "prologue1.png",
    "PROLOGUE_SCENE2": "prologue2.png",
    "CHAPTER_01": "chapter1.png",
    "CHAPTER_02": "chapter2.png",
    "CHAPTER_03": "chapter3.png",
    "CHAPTER_04": "chapter4.png",
    "CHAPTER_05": "chapter5.png",
    "CHAPTER_06": "chapter6.png",
    "CHAPTER_07_5": "chapter7_5.png",
    "CHAPTER_07": "chapter7.png",
    "CHAPTER_08": "chapter8.png",
    "CHAPTER_09": "chapter9.png",
    "CHAPTER_10": "chapter10.png",
    "CHAPTER_11": "chapter11.png",
    "CHAPTER_12": "chapter12.png",
    "CHAPTER_13": "chapter13.png",
    "CHAPTER_14": "chapter14.png",
    "EPILOGUE": "epilogue.png",
}

MAP_IMAGES = [
    "map1-image.png",
    "map2-image.png",
    "map3-image.png",
    "map4-image.png",
    "map5-image.png",
    "map6-image.png",
]

BOOK_TITLES = {
    1: "The Cydonian Oaths",
    2: "The Cauldron of God",
    3: "The Edenic Mandate",
}

# ---------------------------------------------------------------------------
# DOCUMENT SETUP
# ---------------------------------------------------------------------------

def create_document() -> Document:
    doc = Document()

    # Page size: 6" x 9"
    for section in doc.sections:
        section.page_width  = Inches(6)
        section.page_height = Inches(9)
        section.top_margin    = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        # Mirror margins for inside/outside (gutter on binding side)
        section.left_margin  = MARGIN_INSIDE
        section.right_margin = MARGIN_OUTSIDE

    # Remove default blank paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Set default font on Normal style
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(15)

    return doc


# ---------------------------------------------------------------------------
# LOW-LEVEL XML HELPERS
# ---------------------------------------------------------------------------

def add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def add_section_break_odd(doc: Document):
    """Add an odd-page section break (chapter always starts on recto)."""
    new_section = doc.add_section(WD_SECTION.ODD_PAGE)
    new_section.page_width  = Inches(6)
    new_section.page_height = Inches(9)
    new_section.top_margin    = MARGIN_TOP
    new_section.bottom_margin = MARGIN_BOTTOM
    new_section.left_margin  = MARGIN_INSIDE
    new_section.right_margin = MARGIN_OUTSIDE
    return new_section


def set_header_footer(section, author: str, title: str, show_page: bool = True):
    """Verso: author name. Recto: book title. Page number in footer."""
    section.different_first_page_header_footer = True

    # Verso header (left/even pages)
    section.header.is_linked_to_previous = False
    verso = section.header
    if verso.paragraphs:
        p = verso.paragraphs[0]
    else:
        p = verso.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(author.upper())
    run.font.name = BODY_FONT
    run.font.size = Pt(9)

    # Recto header (right/odd pages) — python-docx exposes even_page_header
    # We use the primary header for recto, even_page_header for verso
    # (Word's "different odd/even" must be enabled at doc level)
    enable_odd_even_headers(section._sectPr)

    even_hdr = section.even_page_header
    even_hdr.is_linked_to_previous = False
    if even_hdr.paragraphs:
        ep = even_hdr.paragraphs[0]
    else:
        ep = even_hdr.add_paragraph()
    ep.clear()
    ep.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    erun = ep.add_run(title.upper())
    erun.font.name = BODY_FONT
    erun.font.size = Pt(9)

    if show_page:
        # Default footer (odd pages): page numbers centered
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            fp = footer.paragraphs[0]
        else:
            fp = footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number_field(fp)

        # Even-page footer: page numbers centered
        even_footer = section.even_page_footer
        even_footer.is_linked_to_previous = False
        if even_footer.paragraphs:
            efp = even_footer.paragraphs[0]
        else:
            efp = even_footer.add_paragraph()
        efp.clear()
        efp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number_field(efp)

        # First-page footer: blank (no page number on chapter openers)
        first_footer = section.first_page_footer
        first_footer.is_linked_to_previous = False
        if first_footer.paragraphs:
            ffp = first_footer.paragraphs[0]
        else:
            ffp = first_footer.add_paragraph()
        ffp.clear()


def enable_odd_even_headers(sectPr):
    """Enable different odd/even page headers at section level."""
    titlePg = sectPr.find(qn("w:titlePg"))
    if titlePg is None:
        titlePg = OxmlElement("w:titlePg")
        sectPr.append(titlePg)


def add_page_number_field(paragraph):
    """Insert a PAGE field into a paragraph."""
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")   # 9pt
    rpr.append(sz)
    r.append(rpr)
    fld.append(r)
    run._r.append(fld)


# ---------------------------------------------------------------------------
# INLINE MARKDOWN → RUNS
# ---------------------------------------------------------------------------

def apply_inline(para, text: str, base_bold=False, base_italic=False):
    """
    Parse inline markdown in `text` and add styled runs to `para`.
    Handles: **bold**, *italic*, ***bold-italic***, em-dashes, smart quotes.
    """
    text = text.replace("&nbsp;", " ")
    # Normalise smart quotes (already present in source, preserve them)
    # Convert bare -- to em-dash if missed
    text = re.sub(r"(?<![─—])-{2}(?![─—])", "—", text)

    # Pattern: ***bold-italic***, **bold**, *italic*, plain
    token_re = re.compile(r"(\*{3}.+?\*{3}|\*{2}.+?\*{2}|\*.+?\*)")
    parts = token_re.split(text)

    for part in parts:
        if not part:
            continue
        bold   = base_bold
        italic = base_italic
        content = part

        if part.startswith("***") and part.endswith("***"):
            bold   = True
            italic = True
            content = part[3:-3]
        elif part.startswith("**") and part.endswith("**"):
            bold    = True
            content = part[2:-2]
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            italic  = True
            content = part[1:-1]

        run = para.add_run(content)
        run.font.name   = BODY_FONT
        run.font.size   = BODY_SIZE
        run.bold        = bold
        run.italic      = italic


# Global bookmark ID counter — guarantees unique w:id for every bookmark in the document
_bookmark_id_counter = 0

def _next_bookmark_id() -> str:
    global _bookmark_id_counter
    _bookmark_id_counter += 1
    return str(_bookmark_id_counter)


# ---------------------------------------------------------------------------
# PARAGRAPH FACTORIES
# ---------------------------------------------------------------------------

def add_body_paragraph(doc: Document, text: str, first_para: bool = False) -> None:
    """Standard body paragraph with 0.3" first-line indent (suppressed for first para)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before        = Pt(0)
    p.paragraph_format.space_after         = Pt(0)
    p.paragraph_format.line_spacing_rule   = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing        = Pt(15)
    p.paragraph_format.alignment           = WD_ALIGN_PARAGRAPH.JUSTIFY
    if not first_para:
        p.paragraph_format.first_line_indent = Inches(0.3)
    apply_inline(p, text)


def add_centered(doc: Document, text: str, size: int = 11, bold: bool = False, italic: bool = False, space_before: float = 0, space_after: float = 0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before        = Pt(space_before)
    p.paragraph_format.space_after         = Pt(space_after)
    p.paragraph_format.line_spacing_rule   = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing        = Pt(max(15, size + 4))
    apply_inline(p, text, base_bold=bold, base_italic=italic)
    # Override font size on all runs
    for run in p.runs:
        run.font.size = Pt(size)


def _make_bookmark_id(text: str) -> str:
    """Turn heading text into a stable, legal Word bookmark name (max 40 chars)."""
    # Normalize to uppercase, strip to alphanumeric + underscores
    clean = re.sub(r"[^A-Za-z0-9]+", "_", text.upper()).strip("_")
    name = f"_TOC_{clean}"
    if len(name) > 40:
        name = name[:40]
    return name


def _add_bookmark(paragraph, name: str):
    """Wrap the entire run content of a paragraph in a bookmark."""
    p_elem = paragraph._p
    bs = OxmlElement("w:bookmarkStart")
    bm_id = _next_bookmark_id()
    bs.set(qn("w:id"), bm_id)
    bs.set(qn("w:name"), name)
    # Insert bookmarkStart before first run
    first_r = p_elem.find(qn("w:r"))
    if first_r is not None:
        p_elem.insert(list(p_elem).index(first_r), bs)
    else:
        p_elem.append(bs)
    # Append bookmarkEnd at the end
    be = OxmlElement("w:bookmarkEnd")
    be.set(qn("w:id"), bm_id)
    p_elem.append(be)


def _add_pageref_field(paragraph, bookmark_name: str):
    """Insert a PAGEREF field code into a paragraph."""
    run = paragraph.add_run()
    r = run._r
    # fldChar begin
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    r.append(fc_begin)
    # instrText
    r2 = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")  # 11pt
    rpr.append(sz)
    r2.append(rpr)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark_name} \\h "
    r2.append(instr)
    paragraph._p.append(r2)
    # fldChar separate
    r3 = OxmlElement("w:r")
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r3.append(fc_sep)
    paragraph._p.append(r3)
    # Placeholder text
    r4 = OxmlElement("w:r")
    rpr4 = copy.deepcopy(rpr)
    r4.append(rpr4)
    t4 = OxmlElement("w:t")
    t4.text = "000"
    r4.append(t4)
    paragraph._p.append(r4)
    # fldChar end
    r5 = OxmlElement("w:r")
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r5.append(fc_end)
    paragraph._p.append(r5)


def add_chapter_heading(doc: Document, chapter_line: str, subtitle_line: str = "") -> None:
    """Chapter title block: 2" top space, large centered heading, optional subtitle."""
    # Top space (~2" from top of text area)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after  = Pt(0)
    spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    spacer.paragraph_format.line_spacing = Inches(2)
    spacer.add_run("")

    # Clean the heading text
    heading_text = re.sub(r"^#+\s*", "", chapter_line).strip()
    # Remove inline bold/italic markers for heading (it's all bold anyway)
    heading_text = re.sub(r"\*+", "", heading_text)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after  = Pt(6)
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h.paragraph_format.line_spacing = Pt(26)
    run = h.add_run(heading_text)
    run.font.name  = HEADING_FONT
    run.font.size  = Pt(20)
    run.bold       = True

    # Add bookmark for TOC PAGEREF — only ONE bookmark per heading to avoid
    # nested bookmarks that KDP's converter cannot resolve
    if subtitle_line:
        sub_text = re.sub(r"^#+\s*", "", subtitle_line).strip()
        sub_text = re.sub(r"\*+", "", sub_text)
        bm_name = _make_bookmark_id(f"{heading_text} {sub_text}")
    else:
        bm_name = _make_bookmark_id(heading_text)
    _add_bookmark(h, bm_name)

    if subtitle_line:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.space_after  = Pt(36)
        run2 = s.add_run(sub_text)
        run2.font.name   = HEADING_FONT
        run2.font.size   = Pt(12)
        run2.italic      = True
    else:
        # Still add bottom space
        gap = doc.add_paragraph()
        gap.paragraph_format.space_before = Pt(0)
        gap.paragraph_format.space_after  = Pt(30)
        gap.add_run("")


def add_scene_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(12)
    run = p.add_run(SCENE_BREAK)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE


def add_blank_line(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(15)
    p.add_run("")


# ---------------------------------------------------------------------------
# IMAGE EMBEDDING
# ---------------------------------------------------------------------------

def get_chapter_image(chapter_path: Path, image_dir: Path | None) -> Path | None:
    """Look up the chapter art image for a given chapter file by naming convention."""
    if not image_dir or not image_dir.exists():
        return None
    stem = chapter_path.stem.upper()
    for pattern, img_name in CHAPTER_IMAGE_MAP.items():
        if stem.startswith(pattern):
            img_path = image_dir / img_name
            if img_path.exists():
                return img_path
    return None


def _ensure_chapter_image_style(doc: Document):
    """Create the 'Chapter Image' style matching Book 1's proven definition."""
    style_name = "Chapter Image"
    if style_name not in [s.name for s in doc.styles]:
        img_style = doc.styles.add_style(style_name, 1)  # PARAGRAPH
        img_style.base_style = doc.styles["Normal"]
        img_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        img_style.paragraph_format.line_spacing = 1.0
        img_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_style.paragraph_format.space_before = Pt(0)
        img_style.paragraph_format.space_after = Pt(36)
        img_style.paragraph_format.first_line_indent = Inches(0)
        rpr = img_style.element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            img_style.element.append(rpr)
        no_proof = OxmlElement("w:noProof")
        rpr.append(no_proof)
    return doc.styles[style_name]


def _convert_inline_to_floating(paragraph):
    """Convert the first inline image in a paragraph to a floating anchor.

    Uses 'top and bottom' wrapping, centered horizontally and vertically
    within the margin area.  Floating images are never clipped by
    paragraph line-spacing — this is the nuclear fix for the chapter-art
    display bug.
    """
    for drawing in paragraph._p.findall('.//' + qn('w:drawing')):
        inline = drawing.find(qn('wp:inline'))
        if inline is None:
            continue

        anchor = OxmlElement('wp:anchor')
        anchor.set('distT', '0')
        anchor.set('distB', '0')
        anchor.set('distL', '0')
        anchor.set('distR', '0')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '251658240')
        anchor.set('behindDoc', '0')
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '0')

        # Required: simplePos
        sp = OxmlElement('wp:simplePos')
        sp.set('x', '0')
        sp.set('y', '0')
        anchor.append(sp)

        # Horizontal: centered on margin
        pos_h = OxmlElement('wp:positionH')
        pos_h.set('relativeFrom', 'margin')
        align_h = OxmlElement('wp:align')
        align_h.text = 'center'
        pos_h.append(align_h)
        anchor.append(pos_h)

        # Vertical: centered on margin
        pos_v = OxmlElement('wp:positionV')
        pos_v.set('relativeFrom', 'margin')
        align_v = OxmlElement('wp:align')
        align_v.text = 'center'
        pos_v.append(align_v)
        anchor.append(pos_v)

        # Move children from inline → anchor (extent, effectExtent, docPr, etc.)
        for child in list(inline):
            inline.remove(child)
            anchor.append(child)

        # Insert wrapTopAndBottom before docPr
        wrap = OxmlElement('wp:wrapTopAndBottom')
        doc_pr = anchor.find(qn('wp:docPr'))
        if doc_pr is not None:
            anchor.insert(list(anchor).index(doc_pr), wrap)
        else:
            anchor.append(wrap)

        # Replace inline with anchor
        drawing.remove(inline)
        drawing.append(anchor)
        break  # only convert the first image


def add_chapter_image(doc: Document, image_path: Path,
                      max_width: float = 4.5, max_height: float = 6.5,
                      page_break: bool = False, floating: bool = False) -> None:
    """Embed a centered chapter art image.

    If floating=True, converts the inline picture to a wp:anchor element
    so it is never clipped by paragraph line-spacing.  Maps should keep
    floating=False (inline images work fine without section breaks).
    """
    style = _ensure_chapter_image_style(doc)

    picture = doc.add_picture(str(image_path), width=Inches(max_width))
    if picture.height > Inches(max_height):
        ratio = picture.width / picture.height
        picture.height = Inches(max_height)
        picture.width = int(Inches(max_height) * ratio)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.style = style
    if page_break:
        last_paragraph.paragraph_format.page_break_before = True
    if floating:
        _convert_inline_to_floating(last_paragraph)


# ---------------------------------------------------------------------------
# TABLE RENDERER
# ---------------------------------------------------------------------------

def render_table(doc: Document, rows: list[str]) -> None:
    """Convert a markdown pipe table into a Word table."""
    parsed = []
    for row in rows:
        # Strip leading/trailing pipes and split
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    if not parsed:
        return

    # Remove separator row (contains only dashes and colons)
    parsed = [r for r in parsed if not all(re.match(r"^[-: ]+$", c) for c in r)]

    if not parsed:
        return

    col_count = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row in enumerate(parsed):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= col_count:
                break
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            # Header row: bold
            if r_idx == 0:
                run = p.add_run(re.sub(r"\*+", "", cell_text))
                run.font.name = BODY_FONT
                run.font.size = Pt(9)
                run.bold = True
            else:
                apply_inline(p, cell_text)
                for run in p.runs:
                    run.font.size = Pt(9)

    # Space after table
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(6)
    gap.paragraph_format.space_after  = Pt(6)
    gap.add_run("")


# ---------------------------------------------------------------------------
# FRONT MATTER BUILDER
# ---------------------------------------------------------------------------

def build_front_matter(doc: Document, path: Path,
                       book_number: int = 1, book_title: str = "The Cydonian Oaths") -> None:
    """
    Parse a FRONT_MATTER.md and render each PAGE section
    as properly formatted front matter pages with page breaks.
    """
    book_word = {1: "One", 2: "Two", 3: "Three", 4: "Four", 5: "Five"}.get(book_number, str(book_number))
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()

    # Skip the meta header lines (title of the draft file, typesetter instructions)
    # We'll parse by PAGE markers
    current_page = None
    page_content: dict[str, list[str]] = {}
    page_order = []

    for line in lines:
        page_match = re.match(r"^## PAGE \d+: (.+)$", line)
        if page_match:
            current_page = page_match.group(1).strip()
            page_order.append(current_page)
            page_content[current_page] = []
        elif current_page is not None:
            # Stop collecting if we hit the typesetting notes section
            if line.startswith("# TYPESETTING NOTES"):
                break
            page_content[current_page].append(line)

    def clean_lines(lines_list):
        """Strip leading/trailing blank lines from a page's content."""
        content = [l for l in lines_list if not re.match(r"^---$", l)]
        # Remove bracketed instructions
        content = [l for l in content if not re.match(r"^\*\[.*\]\*$", l.strip())]
        # Strip &nbsp; standalone lines — replace with blank
        content = ["" if l.strip() == "&nbsp;" else l for l in content]
        # Remove leading blank lines
        while content and content[0].strip() == "":
            content.pop(0)
        # Remove trailing blank lines
        while content and content[-1].strip() == "":
            content.pop()
        return content

    # ---- Render each page ----

    for i, page_name in enumerate(page_order):
        content = clean_lines(page_content.get(page_name, []))

        # PAGE BREAK before every page (except the very first)
        if i > 0:
            add_page_break(doc)

        if "SERIES HALF-TITLE" in page_name:
            # Vertical centering approximation: large top space
            for _ in range(8):
                add_blank_line(doc)
            add_centered(doc, "THE NEPHILIM CHRONICLES", size=18, bold=True)
            add_blank_line(doc)
            add_centered(doc, f"Book {book_word}", size=14, italic=True)

        elif "ALSO BY" in page_name:
            add_centered(doc, "Also by Kerman Gild", size=13, bold=True, space_before=24, space_after=12)
            in_series = None
            for line in content:
                stripped = line.strip()
                if not stripped:
                    add_blank_line(doc)
                    continue
                # Series headers (*italic*)
                if re.match(r"^\*[^*].+[^*]\*$", stripped):
                    series_title = stripped.strip("*")
                    add_centered(doc, series_title, size=11, italic=True, space_before=8)
                    in_series = series_title
                elif stripped.startswith("**") and stripped.endswith("**"):
                    add_centered(doc, stripped.strip("*"), size=11, bold=True)
                elif stripped:
                    add_centered(doc, stripped, size=11)

        elif "FULL TITLE PAGE" in page_name:
            for _ in range(6):
                add_blank_line(doc)
            add_centered(doc, "THE", size=14, bold=True)
            add_centered(doc, "NEPHILIM", size=14, bold=True)
            add_centered(doc, "CHRONICLES", size=14, bold=True)
            add_blank_line(doc)
            add_centered(doc, book_title.upper(), size=26, bold=True, space_before=6, space_after=6)
            add_blank_line(doc)
            add_centered(doc, f"Book {book_word}", size=13, italic=True)
            for _ in range(4):
                add_blank_line(doc)
            add_centered(doc, "KERMAN GILD", size=14, bold=True)
            add_blank_line(doc)
            add_centered(doc, "Kerman Gild Publishing", size=11)
            add_centered(doc, "Auckland, New Zealand", size=11)

        elif "COPYRIGHT" in page_name:
            add_blank_line(doc)
            for line in content:
                stripped = line.strip()
                if not stripped:
                    add_blank_line(doc)
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                p.paragraph_format.line_spacing = Pt(14)
                apply_inline(p, stripped)
                for run in p.runs:
                    run.font.size = Pt(9)

        elif "DEDICATION" in page_name:
            for _ in range(7):
                add_blank_line(doc)
            for line in content:
                stripped = line.strip()
                if not stripped:
                    add_blank_line(doc)
                    continue
                # Each dedication line is italic
                is_italic = stripped.startswith("*") and stripped.endswith("*")
                clean = stripped.strip("*")
                add_centered(doc, clean, size=11, italic=True)

        elif "BLANK" in page_name:
            # Intentionally blank — just the page break already added is enough
            add_blank_line(doc)

        elif "EPIGRAPH" in page_name:
            for _ in range(5):
                add_blank_line(doc)
            epigraph_text = []
            attribution = ""
            for line in content:
                stripped = line.strip()
                if stripped.startswith("—") or stripped.startswith("\u2014"):
                    attribution = stripped
                    # Render the accumulated epigraph block
                    if epigraph_text:
                        for et in epigraph_text:
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after  = Pt(0)
                            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                            p.paragraph_format.line_spacing = Pt(16)
                            p.paragraph_format.left_indent  = Inches(0.75)
                            p.paragraph_format.right_indent = Inches(0.75)
                            apply_inline(p, et.strip().strip("*"), base_italic=True)
                            for r in p.runs:
                                r.font.size = Pt(10)
                        epigraph_text = []
                    # Attribution
                    attr_p = doc.add_paragraph()
                    attr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    attr_p.paragraph_format.space_before = Pt(2)
                    attr_p.paragraph_format.space_after  = Pt(18)
                    attr_p.paragraph_format.right_indent = Inches(0.75)
                    run = attr_p.add_run(attribution.lstrip("*").rstrip("*"))
                    run.font.name = BODY_FONT
                    run.font.size = Pt(9)
                    attribution = ""
                elif stripped:
                    epigraph_text.append(stripped)

        elif "TABLE OF CONTENTS" in page_name:
            add_centered(doc, "CONTENTS", size=16, bold=True, space_before=36, space_after=18)
            for line in content:
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith("**") and stripped.endswith("**"):
                    # Bold heading ("CONTENTS" — skip, already rendered)
                    continue
                # Build the bookmark name this entry points to
                # Match the heading text portion after the em-dash
                bm_name = _make_bookmark_id(stripped)
                # TOC entry with right-aligned tab + dot leader + PAGEREF
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                p.paragraph_format.line_spacing = Pt(16)
                p.paragraph_format.left_indent = Inches(0.75)
                # Add a right tab stop with dot leader at the right margin
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(4.625), WD_ALIGN_PARAGRAPH.RIGHT, 2)  # 2 = dots leader
                # Entry text
                run = p.add_run(stripped)
                run.font.name = BODY_FONT
                run.font.size = Pt(11)
                # Tab character to jump to dot leader
                tab_run = p.add_run("\t")
                tab_run.font.name = BODY_FONT
                tab_run.font.size = Pt(11)
                # PAGEREF field — updates to real page numbers via Ctrl+A, F9 in Word
                _add_pageref_field(p, bm_name)

        elif "PREFACE" in page_name:
            add_centered(doc, "A Word Before the Story", size=14, bold=True, space_before=36, space_after=18)
            first = True
            for line in content:
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith("#"):
                    continue  # already rendered as heading above
                if stripped.startswith("*—") or stripped == "— Kerman Gild" or re.match(r"^\*—\s*Kerman Gild\*$", stripped):
                    add_blank_line(doc)
                    add_centered(doc, "— Kerman Gild", size=11, italic=True)
                    continue
                add_body_paragraph(doc, stripped, first_para=first)
                first = False


# ---------------------------------------------------------------------------
# CHAPTER / BODY FILE RENDERER
# ---------------------------------------------------------------------------

def is_table_row(line: str) -> bool:
    return bool(re.match(r"^\s*\|.+\|", line))


def render_body_file(doc: Document, path: Path, is_first_file: bool = False,
                     image_dir: Path | None = None,
                     book_title: str = "The Cydonian Oaths") -> None:
    """
    Render a chapter, prologue, epilogue, or appendix file.
    Inserts an odd-page section break before it (so chapters start on recto).
    Embeds chapter art image after heading if available in image_dir.
    """
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()

    # Detect title (#) and subtitle (##) lines
    title_line    = ""
    subtitle_line = ""
    body_start    = 0

    for idx, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# ") and not title_line:
            title_line = stripped
            body_start = idx + 1
        elif stripped.startswith("## ") and title_line and not subtitle_line:
            subtitle_line = stripped
            body_start = idx + 1
        else:
            break

    # Chapter art on its own page before the heading (if available)
    chapter_image = get_chapter_image(path, image_dir)
    if chapter_image:
        add_chapter_image(doc, chapter_image, page_break=True, floating=True)
        print(f"      art: {chapter_image.name}")

    # Section break (odd page) so chapter starts on recto
    section = add_section_break_odd(doc)
    set_header_footer(
        section,
        author="Kerman Gild",
        title=book_title,
        show_page=True,
    )

    # Chapter heading block
    add_chapter_heading(doc, title_line, subtitle_line)

    # Body
    first_para = True
    table_buffer: list[str] = []
    in_table = False
    skip_next_blank = False

    # Lines that are internal writer's notes (character voice, timeline notes etc.)
    NOTE_PATTERNS = [
        r"^#{3,}",              # ### headings (scene notes inside manuscript)
        r"^\*\*\*$",            # bare triple asterisk (not scene break)
        r"^> ",                 # blockquotes (author notes)
        r"^\[.*\]$",            # raw bracket instructions
    ]
    # But we DO want to pass through ### inside appendices as sub-headings
    is_appendix = "APPENDIX" in path.name or "APPENDICES" in path.name

    i = body_start
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        i += 1

        # Skip blank / &nbsp; lines — just mark paragraph boundary
        if not stripped or stripped == "&nbsp;":
            skip_next_blank = False
            first_para = False   # after any blank, next paragraph gets indent
            continue

        # Skip typesetter-note lines (remain in chapter manuscripts)
        if any(re.match(pat, stripped) for pat in NOTE_PATTERNS[:-1]):
            if is_appendix and stripped.startswith("###"):
                # Render as sub-heading
                heading_text = re.sub(r"^#{1,6}\s*", "", stripped)
                add_centered(doc, heading_text, size=11, bold=True, space_before=12, space_after=4)
                first_para = True
                continue
            # Skip writer's note lines in chapters
            continue

        # Markdown headings inside body (## scene location tags etc.)
        if stripped.startswith("## "):
            heading_text = re.sub(r"^#+\s*", "", stripped)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after  = Pt(12)
            run = p.add_run(heading_text)
            run.font.name   = HEADING_FONT
            run.font.size   = Pt(12)
            run.italic      = True
            first_para = True
            continue

        if stripped.startswith("# "):
            # Second-level title inside appendices (Appendix A, B…)
            heading_text = re.sub(r"^#+\s*", "", stripped)
            add_centered(doc, heading_text, size=16, bold=True, space_before=24, space_after=8)
            first_para = True
            continue

        # Scene break: --- or ***
        if stripped in ("---", "* * *", "***") or stripped == SCENE_BREAK:
            add_scene_break(doc)
            first_para = True
            continue

        # Table detection
        if is_table_row(stripped):
            table_buffer.append(stripped)
            # Peek ahead — collect whole table
            while i < len(lines) and is_table_row(lines[i].strip()):
                table_buffer.append(lines[i].strip())
                i += 1
            render_table(doc, table_buffer)
            table_buffer = []
            first_para = True
            continue

        # Horizontal rule in appendices (---) → scene break
        if stripped == "---" and is_appendix:
            add_scene_break(doc)
            first_para = True
            continue

        # Code fences → skip (shouldn't appear in final manuscript)
        if stripped.startswith("```"):
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1  # skip closing fence
            continue

        # Empyreal Register speech (all caps bold within **) — already handled by inline parser
        # Dialogue tables / character notes at end of chapters — skip
        if stripped.startswith("| ") and not is_appendix:
            # Skip manuscript tables that are author notes
            while i < len(lines) and lines[i].strip().startswith("|"):
                i += 1
            first_para = True
            continue

        # Normal body paragraph
        add_body_paragraph(doc, stripped, first_para=first_para)
        first_para = False


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    book = BOOK

    # Resolve per-book configuration
    if book == 2:
        manuscript_dir = MANUSCRIPT_DIR_BOOK2
        output_file = OUTPUT_FILE_BOOK2
        source_files = SOURCE_FILES_BOOK2
        image_dir = IMAGE_DIR_BOOK2
        book_title = BOOK_TITLES[2]
        has_front_matter = True
    elif book == 3:
        manuscript_dir = MANUSCRIPT_DIR_BOOK3
        output_file = OUTPUT_FILE_BOOK3
        source_files = SOURCE_FILES_BOOK3
        image_dir = IMAGE_DIR_BOOK3
        book_title = BOOK_TITLES[3]
        has_front_matter = False  # Book 3 uses chapter files only (no front matter yet)
    else:
        manuscript_dir = MANUSCRIPT_DIR
        output_file = OUTPUT_FILE
        source_files = SOURCE_FILES
        image_dir = None
        book_title = BOOK_TITLES[1]
        has_front_matter = True

    # Env var overrides
    if os.environ.get("KDP_MANUSCRIPT_DIR"):
        manuscript_dir = Path(os.environ["KDP_MANUSCRIPT_DIR"])
    if os.environ.get("KDP_OUTPUT_FILE"):
        output_file = Path(os.environ["KDP_OUTPUT_FILE"])

    print(f"Building manuscript: Book {book} — {book_title}")
    doc = create_document()

    # Global odd/even headers + mirror margins (gutter on binding side for all pages)
    settings = doc.settings.element
    evenOdd = OxmlElement("w:evenAndOddHeaders")
    settings.append(evenOdd)
    mirror = OxmlElement("w:mirrorMargins")
    settings.append(mirror)

    # ── Front Matter ─────────────────────────────────────────────────────────
    if has_front_matter:
        front_matter_file = source_files[0]
        if front_matter_file.exists():
            print("  → Front matter")
            build_front_matter(doc, front_matter_file,
                               book_number=book, book_title=book_title)
        body_files = source_files[1:]
    else:
        body_files = source_files

    # ── Body files ───────────────────────────────────────────────────────────
    for idx, filepath in enumerate(body_files):
        if not filepath.exists():
            print(f"  ⚠  Missing: {filepath.name} — skipping")
            continue
        print(f"  → {filepath.name}")
        render_body_file(doc, filepath, is_first_file=(idx == 0),
                         image_dir=image_dir, book_title=book_title)

    # ── Maps section (Book 2) ────────────────────────────────────────────────
    if book == 2 and image_dir:
        maps_found = [image_dir / m for m in MAP_IMAGES if (image_dir / m).exists()]
        if maps_found:
            section = add_section_break_odd(doc)
            set_header_footer(section, author="Kerman Gild",
                              title=book_title, show_page=True)
            add_chapter_heading(doc, "# MAPS")
            for map_path in maps_found:
                add_chapter_image(doc, map_path, max_width=4.5, max_height=7.0)
                print(f"    map: {map_path.name}")

    # ── Save ─────────────────────────────────────────────────────────────────
    doc.save(output_file)
    print(f"\n✓  Saved: {output_file}")
    print(f"   Word count estimate: check document properties after opening.")


if __name__ == "__main__":
    main()
